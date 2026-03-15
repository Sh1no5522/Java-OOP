import os
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

EXCEL_FILE = r"Batch 2 Group 12-M.Aidan.xlsx"
MERGED_DOC = r"merged_document.docx"
OUTPUT_DOC = r"Final_Report_MAidan.docx"


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_para(doc, text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)
    return p


def add_table(doc, dataframe, title=None):
    if title:
        add_para(doc, title)

    table = doc.add_table(rows=1, cols=len(dataframe.columns))
    hdr = table.rows[0].cells
    for i, col in enumerate(dataframe.columns):
        hdr[i].text = str(col)

    for _, row in dataframe.iterrows():
        cells = table.add_row().cells
        for i, item in enumerate(row):
            cells[i].text = str(item)

    add_para(doc, "")  # space


df = pd.read_excel(EXCEL_FILE)

# If your first column is respondent ID, drop it
if df.columns[0].lower() in ["id", "respondent", "number"]:
    df = df.iloc[:, 1:]


desc_stats = df.describe().T


merged = Document(MERGED_DOC)
merged_text = "\n".join([p.text for p in merged.paragraphs])

# Split by sections (you can adjust keywords)
sections = {}
current = None
for line in merged_text.split("\n"):
    line_strip = line.strip()

    if line_strip.startswith("Interpretation"):
        current = "Interpretation"
        sections[current] = ""
    elif line_strip.startswith("Findings"):
        current = "Findings"
        sections[current] = ""
    elif line_strip.startswith("Conclusion"):
        current = "Conclusion"
        sections[current] = ""
    elif line_strip.startswith("Recommendation"):
        current = "Recommendation"
        sections[current] = ""

    if current:
        sections[current] += line + "\n"

# -----------------------------
# STEP 4: CREATE FINAL REPORT
# -----------------------------
doc = Document()

# TITLE PAGE
add_heading(doc, "FINAL RESEARCH REPORT", level=1)
add_para(doc, "Prepared by: Mustafin Aidan")
add_para(doc, "Course: Statistics / Research Methods")
add_para(doc, " ")

# -----------------------------
# ABSTRACT (SHORT)
# -----------------------------
add_heading(doc, "Abstract", 1)
add_para(doc,
         "This study investigates the impact of ERP implementation on supply chain "
         "visibility among SMEs in Kazakhstan. The research evaluates system usage, "
         "data accuracy, integration quality, and operational improvement. Using "
         "statistical analysis, descriptive patterns, reliability checks, and hypothesis "
         "testing, the study provides evidence on how ERP adoption influences "
         "performance outcomes.")

# -----------------------------
# INTRODUCTION
# -----------------------------
add_heading(doc, "1. Introduction", 1)
add_para(doc,
         "The purpose of this study is to examine the role of Enterprise Resource "
         "Planning (ERP) systems in enhancing supply chain visibility within SMEs "
         "in Kazakhstan. Supply chain visibility refers to real-time information "
         "sharing, transparency, and operational coordination between business units. "
         "ERP systems play a critical role by integrating data across functions such "
         "as procurement, inventory, logistics, and finance. This research aims to "
         "evaluate several ERP-related constructs and determine their effects on "
         "overall organizational performance.")

# -----------------------------
# METHODOLOGY
# -----------------------------
add_heading(doc, "2. Research Methodology", 1)
add_para(doc,
         "A quantitative research design was adopted. Data were collected using a "
         "structured questionnaire based on Likert-scale items. Respondents were "
         "drawn from SMEs operating in Kazakhstan. Data analysis included descriptive "
         "statistics, reliability testing, regression analysis, and hypothesis testing.")

# -----------------------------
# DESCRIPTIVE STATISTICS
# -----------------------------
add_heading(doc, "3. Descriptive Statistics", 1)
add_table(doc, desc_stats, title="Table 1: Descriptive Statistics")

# -----------------------------
# INTERPRETATION SECTION (FROM MERGED FILE)
# -----------------------------
if "Interpretation" in sections:
    add_heading(doc, "4. Interpretation of Findings", 1)
    add_para(doc, sections["Interpretation"])

# -----------------------------
# FINDINGS
# -----------------------------
if "Findings" in sections:
    add_heading(doc, "5. Findings", 1)
    add_para(doc, sections["Findings"])

# -----------------------------
# CONCLUSION
# -----------------------------
if "Conclusion" in sections:
    add_heading(doc, "6. Conclusion", 1)
    add_para(doc, sections["Conclusion"])

# -----------------------------
# RECOMMENDATION
# -----------------------------
if "Recommendation" in sections:
    add_heading(doc, "7. Recommendations", 1)
    add_para(doc, sections["Recommendation"])

# -----------------------------
# SAVE FINAL REPORT
# -----------------------------
doc.save(OUTPUT_DOC)

print("Final report generated successfully →", OUTPUT_DOC)
